# ingestion/loader.py
import os
from ingestion.schema import Document
from docx import Document as DocxDocument


def load_txt_files(data_dir: str):
    documents = []
    for fname in os.listdir(data_dir):
        if fname.endswith(".txt"):
            path = os.path.join(data_dir, fname)
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()

        elif fname.endswith(".docx"):
            path = os.path.join(data_dir, fname)
            docx = DocxDocument(path)
            text = "\n".join([p.text for p in docx.paragraphs])

        else:
            continue

        doc = Document(
            doc_id=fname,
            title=fname.replace(".txt", "").replace(".docx", ""),
            text=text,
            source="internal_kb",
            metadata={"filename": fname}
        )
        documents.append(doc)

    return documents
